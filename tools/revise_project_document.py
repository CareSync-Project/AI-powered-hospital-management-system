from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parents[1] / "Smart_Hospital_Management_System_REVISED.docx"

def paras(doc, values):
    for value in values:
        p = doc.add_paragraph(value)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def bullets(doc, values, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for value in values:
        doc.add_paragraph(value, style=style)

def heading(doc, value, level=1):
    doc.add_heading(value, level=level)

def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(headers):
        t.rows[0].cells[i].text = value
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "007F73")
        t.rows[0].cells[i]._tc.get_or_add_tcPr().append(shd)
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()

def chapter(doc, title):
    doc.add_page_break()
    heading(doc, title, 1)

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(.8)
sec.left_margin = sec.right_margin = Inches(.9)
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(12)
doc.styles["Normal"].paragraph_format.line_spacing = 1.5
for name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
    doc.styles[name].font.name = "Times New Roman"
for name in ["Heading 1", "Heading 2"]:
    doc.styles[name].font.color.rgb = RGBColor(0, 95, 87)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(100)
r = p.add_run("CARESYNC\n")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(0,127,115)
r = p.add_run("AI-POWERED HOSPITAL MANAGEMENT SYSTEM\n")
r.bold = True; r.font.size = Pt(20)
r = p.add_run("WITH INTELLIGENT APPOINTMENT AND CLINICAL WORKFLOW SUPPORT")
r.bold = True; r.font.size = Pt(16)
p = doc.add_paragraph("\nRevised Project Documentation\nAligned with the implemented CareSync application\nSeptember 2026")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

chapter(doc, "DOCUMENT REVISION NOTE")
paras(doc, [
    "This edition replaces speculative statements in the original report with a description of functionality demonstrably implemented in the repository. Completed features, limitations, and future work are explicitly separated. CareSync is a single-hospital academic prototype configured for the fictional CareSync Hospital; it is not a certified electronic health record, diagnostic device, or production clinical system.",
    "The implementation consists of a React web client, an Express REST API, PostgreSQL accessed through Prisma, and a private FastAPI machine-learning service. The supported roles are Administrator, Doctor, Nurse, and Patient. Docker Compose packages the services for local execution.",
])
heading(doc, "Executive Summary", 2)
paras(doc, [
    "CareSync coordinates patient registration, appointment discovery and booking, staff scheduling, nurse assignment, check-in, nurse-recorded vital signs, triage, waiting queues, doctor consultation, patient clinical summaries, notifications, announcements, reporting, and preliminary symptom assessment. Authorization is enforced by the API rather than trusted to browser state.",
    "The term intelligent refers to bounded decision support. Appointment suggestions use real generated slots and explicit scheduling constraints. Symptom assessment screens emergencies first, then uses a measured seven-category classifier when available and a rule-based fallback otherwise. Outputs are advisory and cannot set nurse urgency or doctor diagnosis. The Care Assistant uses deterministic intents and controlled service functions; no external large language model is configured.",
])
table(doc, ["Area", "Actual implementation"], [
    ("Security", "PostgreSQL identities, bcrypt, JWT access tokens, rotating refresh sessions, RBAC and audits"),
    ("Patient", "Registration, profile, cards, booking, symptom check, read-only nurse vitals, notifications and PWA"),
    ("Nurse", "Assigned-patient worklist, check-in, assigned-nurse vital entry, triage and queue placement"),
    ("Doctor", "Assigned queue, clinical context, consultation, treatment, completion and follow-up"),
    ("Admin", "Hospital, departments, staff, schedules, assignments, announcements, reports and analytics"),
    ("AI", "Red-flag rules, logistic-regression classifier, qualitative output, fallback and safe assistant"),
    ("Runtime", "Nginx/React client, Express API, FastAPI inference, PostgreSQL and Docker Compose"),
])

chapter(doc, "CHAPTER ONE: INTRODUCTION")
heading(doc, "1.1 Background of the Study", 2)
paras(doc, [
    "Hospitals coordinate patient identity, provider availability, appointments, intake, queues, consultation, communication, and reporting. Fragmented or paper-based processes can duplicate work, obscure a patient's progress, and create avoidable scheduling conflicts. Digitizing a calendar alone is insufficient: a dependable system must enforce clinic schedules, staff assignments, slot capacity, record ownership, workflow status, and clinical authority.",
    "CareSync was developed as a full-stack academic prototype connecting appointment administration with a controlled outpatient workflow. It centralizes records in PostgreSQL, supplies role-specific web interfaces, and adds carefully bounded artificial-intelligence support for preliminary symptom routing. It demonstrates technical feasibility and workflow integrity, not measured improvement in a real hospital.",
])
heading(doc, "1.2 Statement of the Problem", 2)
paras(doc, [
    "Disconnected appointment processes can cause double booking, uneven workload, poor communication, and weak accountability. Systems that allow patients or unrelated staff to alter clinical information introduce additional safety risks. The project addresses coordination and trust by booking only genuine availability, restricting clinical actions to assigned staff, protecting patient ownership, enforcing valid status transitions, and keeping AI subordinate to clinical judgment.",
])
heading(doc, "1.3 Aim", 2)
paras(doc, ["To design, implement, and evaluate a secure full-stack hospital-management prototype that coordinates appointment scheduling and outpatient clinical workflow while providing bounded AI-assisted symptom and routing support."])
heading(doc, "1.4 Objectives", 2)
bullets(doc, [
    "Implement secure patient registration, administrator-created staff accounts, authentication, session rotation, and server-side role authorization.",
    "Manage the hospital, departments, doctors, nurses, recurring schedules, exceptions, assignments, and appointment slots.",
    "Allow patients to discover real availability and book, reschedule, or cancel their own appointments.",
    "Implement check-in, assigned-nurse vital entry and triage, deterministic queues, and assigned-doctor consultation.",
    "Provide notifications, announcements, reports, analytics, audit records, and role-scoped dashboards.",
    "Integrate emergency-first hybrid symptom assessment and a controlled Care Assistant without claiming diagnosis.",
    "Package and test the client, API, AI service, and database with Docker Compose.",
], True)
heading(doc, "1.5 Research Questions", 2)
bullets(doc, [
    "How can role and relationship checks improve the integrity of appointment and clinical data?",
    "How can schedules, exceptions, and slot capacity support reliable patient self-booking?",
    "How can each role receive appropriate functions without exposing unrelated records?",
    "How can AI assist symptom routing while preserving emergency rules and human authority?",
    "Does the implemented prototype satisfy its functional requirements under automated testing?",
], True)
heading(doc, "1.6 Scope", 2)
bullets(doc, [
    "One configured facility and four roles: ADMIN, DOCTOR, NURSE, and PATIENT.",
    "Patient registration, profile, cards, appointments, symptom assessment, notifications, read-only vitals, and finalized summaries.",
    "Staff, department, schedule, exception, slot, nurse-assignment, announcement, report, and analytics management.",
    "Clinical progression from check-in through nurse intake and doctor consultation.",
    "Hybrid preliminary symptom support and confirmation-based appointment handoff.",
])
heading(doc, "1.7 Limitations", 2)
bullets(doc, [
    "Academic prototype; no clinical certification or prospective hospital validation.",
    "No billing, pharmacy, laboratory, radiology, inpatient, insurer, or external-EHR integration.",
    "No exact wait prediction, automatic clinical urgency, autonomous diagnosis, or prescription.",
    "The ML dataset is public, medically inspired, non-Ghanaian, and not clinical evidence.",
    "Email requires deployment configuration; in-app notification is the authoritative application record.",
    "Docker-based local operability is not a production security or scalability certification.",
])

chapter(doc, "CHAPTER TWO: LITERATURE REVIEW")
heading(doc, "2.1 Hospital Information Systems", 2)
paras(doc, [
    "Hospital information systems centralize administrative and clinical information, but success depends on usability, security, interoperability, infrastructure, training, and maintenance. OpenMRS, OpenEMR, Bahmni, GNU Health, and HospitalRun demonstrate broad combinations of records, scheduling, reporting, and low-resource deployment. CareSync is deliberately narrower: it focuses on secure appointment-to-consultation coordination and bounded decision support.",
    "Research on post-implementation challenges in Ghana and on health-information-system obstacles shows why technical functionality alone is insufficient. Organizational adoption, data quality, governance, support, privacy, and infrastructure must be considered alongside software features.",
])
heading(doc, "2.2 Scheduling and Decision Support", 2)
paras(doc, [
    "CareSync separates deterministic scheduling from machine learning. Department sessions, doctor availability, exceptions, maximum patients, and transactional capacity are explicit business rules. Patient recommendations are selected from actual available database slots. Machine learning is used only for preliminary symptom categorization, after emergency screening, and cannot control clinical triage or diagnosis.",
])
heading(doc, "2.3 Security and Research Gap", 2)
paras(doc, [
    "Healthcare access control must combine roles with record ownership, hospital scope, staff assignment, and workflow state. The gap addressed is a testable integration of real scheduling constraints, patient self-service, assigned-nurse intake, deterministic queue ordering, doctor-owned consultation, safe patient feedback, and explicitly limited AI support.",
])
heading(doc, "2.4 Selected References", 2)
paras(doc, [
    "Anian, S. et al. (2025). A multi-country review of the governance of hospital information systems interoperability. East Mediterranean Health Journal, 31(9&10), 581–589.",
    "Fatima, A., & Colomo-Palacios, R. (2018). Security aspects in healthcare information systems: A systematic mapping. Procedia Computer Science, 138, 12–19.",
    "Mensah, N. K. et al. (2023). Electronic health records post-implementation challenges in selected hospitals in southern Ghana. Digital Health, 9.",
    "Tummers, J. et al. (2021). Obstacles and features of health information systems: A systematic literature review. Computers in Biology and Medicine, 137, 104785.",
])

chapter(doc, "CHAPTER THREE: METHODOLOGY AND SYSTEM DESIGN")
heading(doc, "3.1 Development Method", 2)
paras(doc, [
    "The project used incremental, test-supported development. Phases introduced the data foundation, authentication, administration and scheduling, patient booking, clinical workflow, AI support, and review corrections. Requirements were expressed as actor permissions, relationship rules, validation constraints, and legal state transitions. The repository and running services were treated as the source of truth for this report.",
])
heading(doc, "3.2 Architecture", 2)
table(doc, ["Layer", "Technology", "Responsibility"], [
    ("Web client", "React 19, Vite, Router", "Role-specific interfaces, PWA shell and API consumption"),
    ("REST API", "Node.js 20, Express, Zod", "Authentication, authorization, validation and workflow orchestration"),
    ("Persistence", "PostgreSQL, Prisma", "Authoritative relational records and transactions"),
    ("AI service", "FastAPI, scikit-learn", "Private symptom-model inference"),
    ("Deployment", "Docker Compose, Nginx", "Local multi-service operation"),
])
paras(doc, ["The browser calls Express, never FastAPI directly. Express retains authentication, emergency rules, department resolution, persistence, auditing, and fallback responsibility. This prevents model availability from bypassing core safety and access controls."])
heading(doc, "3.3 Roles and Access Boundaries", 2)
table(doc, ["Role", "Permitted work", "Restriction"], [
    ("Patient", "Own profile, cards, booking, assessments, notifications, vitals and summaries", "Vitals are read-only; cannot triage or diagnose"),
    ("Nurse", "Assigned patients, check-in, vitals, triage and waiting transition", "Vital entry requires active appointment assignment"),
    ("Doctor", "Own queue, clinical context, consultation and follow-up", "Cannot access unassigned consultations or enter nurse vitals"),
    ("Admin", "Hospital, departments, staff, schedules, assignments, announcements and reports", "Cannot rewrite clinical diagnosis through reports"),
])
heading(doc, "3.4 Scheduling and Clinical Workflow", 2)
paras(doc, [
    "Recurring clinic and doctor schedules generate deterministic slots. Leave, holidays, unavailability, and custom hours affect generation. Booked slots survive regeneration. Appointment creation and capacity updates share a serializable transaction, protecting against concurrent overbooking.",
    "The normal clinical state sequence is CONFIRMED → CHECKED_IN → TRIAGED → WAITING → IN_CONSULTATION → COMPLETED. Conditional updates prevent arbitrary jumps. Queue priority is Emergency, High, Moderate, Low, and Routine, with time-based tie breaking.",
    "Vitals use an appointment-scoped nurse endpoint. The API derives patient, hospital, appointment, recorder, source=NURSE, and verificationStatus=VERIFIED, and confirms an active nurse assignment. BMI is computed only when height and weight are present. Patients see verified nurse records but cannot submit or edit them.",
])
heading(doc, "3.5 Authentication and Security", 2)
paras(doc, [
    "Public registration is patient-only; administrators create staff. Passwords use bcrypt cost 12. Login supplies a 15-minute JWT access token and a seven-day random refresh token in an HttpOnly cookie. Only the refresh-token hash is stored, and refresh rotates the token. Protected requests enforce role, ownership, hospital membership, assignment, and state rules. Validation and safe serialization prevent database details and credential hashes from reaching clients.",
])
heading(doc, "3.6 AI Design", 2)
paras(doc, [
    "The symptom pipeline normalizes input and applies emergency red-flag rules before inference. The selected logistic-regression artifact maps 23 symptom features to seven broad outpatient categories. Results use qualitative strengths rather than disease probabilities. Low confidence or AI-service failure activates a rule-based fallback.",
    "The Care Assistant recognizes controlled intents for departments, schedules, availability, appointments, card status, symptom checks, and booking handoff. It cannot execute arbitrary SQL, retain unrestricted chat history, reserve a slot without confirmation, diagnose, prescribe, set triage urgency, or modify doctor findings.",
])
heading(doc, "3.7 Principal Data Groups", 2)
bullets(doc, [
    "Identity and organization: users, sessions, role profiles, hospital, departments and staff relationships.",
    "Scheduling: department schedules, doctor schedules, exceptions and appointment slots.",
    "Care: appointments, nurse assignments, vital records, triage records and consultations.",
    "Patient communication: cards, symptom assessments, notifications and announcements.",
    "Governance: audit logs and operational metadata.",
])
heading(doc, "3.8 Testing Strategy", 2)
bullets(doc, [
    "Client linting and production builds.",
    "Vitest integration tests for authentication, authorization, validation, schedules, booking, clinical workflow, AI integration and business rules.",
    "Python tests for feature engineering, model loading, inference, emergency bypass and fallback.",
    "Docker image builds, container status checks and HTTP health checks.",
])

chapter(doc, "CHAPTER FOUR: IMPLEMENTATION, TESTING, AND RESULTS")
heading(doc, "4.1 Implemented Modules", 2)
table(doc, ["Module", "Implemented behavior"], [
    ("Authentication", "Patient signup; staff provisioning; login, refresh, logout, password change and safe profile context"),
    ("Administration", "Hospital settings, departments, staff, assignments, schedules, slot generation, announcements and reporting"),
    ("Patient portal", "Booking, cancellation, rescheduling, profile, cards, symptom assessment, notifications, vitals and summaries"),
    ("Nurse portal", "Assigned patients, check-in, vital entry, triage and queue movement"),
    ("Doctor portal", "Assigned queue, consultation context, draft, completion and follow-up"),
    ("AI/Care Assistant", "Emergency-first hybrid symptom support and controlled information/booking handoff"),
])
heading(doc, "4.2 Patient Appointment Implementation", 2)
paras(doc, [
    "The patient selects a department, date, doctor, and generated slot. The backend validates patient identity, active hospital and department, doctor assignment, optional card ownership, and capacity. Recommendation logic returns an earliest suitable real slot rather than inventing availability. Cancellation is restricted by status; rescheduling atomically releases the old capacity and reserves the new slot.",
])
heading(doc, "4.3 Nurse Vitals and Triage Implementation", 2)
paras(doc, [
    "An administrator assigns a nurse to an appointment. The nurse sees an assigned-patient worklist and may record bounded measurements only while the assignment is active and the appointment is in an allowed clinical state. Recorded values include applicable temperature, blood pressure, pulse, respiratory rate, oxygen saturation, height, weight, blood glucose, pain score, and notes. BMI is derived when possible.",
    "The patient's My Vitals page contains no input form. It retrieves the authenticated patient's history and displays only verified nurse-sourced values. This design makes the nurse the accountable recorder and prevents patient-entered data from appearing as clinician-verified measurements.",
])
heading(doc, "4.4 Consultation Implementation", 2)
paras(doc, [
    "Only the assigned doctor can start a waiting appointment. Start creates or initializes the unique consultation and timestamps the transition. Draft saving does not complete care. Completion requires clinician-entered diagnosis and treatment plan; a follow-up selection requires a future date. Completion updates both consultation and appointment, audits the action, and creates a patient notification. Patients receive a restricted finalized summary, not internal notes or audits.",
])
heading(doc, "4.5 Notifications, Announcements, and Reporting", 2)
paras(doc, [
    "Notifications are owned records and read operations use authenticated identity. Administrators can publish announcements to supported audiences, including patients, and recipients see them on the notification page. Email is attempted only when mail configuration and the relevant delivery path are available; the document does not claim guaranteed email delivery.",
    "Administrative reports query PostgreSQL and support date, department, doctor, nurse, and status filters. Authorized results can be exported as CSV in the browser. Doctors receive ownership-scoped daily and follow-up reporting. Analytics describe stored prototype activity and are not evidence of real-world hospital performance.",
])
heading(doc, "4.6 Machine-Learning Implementation and Results", 2)
paras(doc, [
    "The SympScan public dataset originally contained 96,088 rows and 230 symptom columns. Nineteen thousand and ninety-two selected rows mapped into seven safer categories. Removing duplicate projected feature/target patterns left 1,314 examples. The stratified split used 919 training, 197 validation, and 198 untouched test examples with random seed 42.",
])
table(doc, ["Candidate", "Validation accuracy", "Macro F1", "Weighted F1"], [
    ("Logistic Regression", "0.720812", "0.727733", "0.731753"),
    ("Random Forest", "0.598985", "0.580132", "0.608352"),
    ("Gradient Boosting", "0.659898", "0.622507", "0.671138"),
])
paras(doc, [
    "Logistic Regression was selected. On the untouched test set it achieved accuracy 0.696970, macro F1 0.676433, and weighted F1 0.707478. These values characterize this dataset split only. They do not constitute clinical accuracy, Ghanaian validation, probability calibration, or regulatory evidence.",
])
heading(doc, "4.7 Verification Results", 2)
paras(doc, [
    "The revised assigned-nurse vitals workflow was verified by the focused clinical suite: 21 of 21 tests passed. The suite includes rejection of an unassigned nurse, enforcement of NURSE and VERIFIED metadata, BMI calculation, hospital isolation, technical bounds, valid workflow transitions, queue ordering, assigned-doctor access, concurrent consultation protection, completion notification, and patient-safe summaries.",
    "The client lint and production build completed successfully, with only pre-existing non-blocking warnings. Server syntax checks passed. Updated server and client Docker images built successfully, containers restarted, the API health endpoint responded successfully, and the frontend returned HTTP 200. Wider repository suites cover authentication, validation, schemas, scheduling, patient booking, symptom assessment, rate limiting, Care Assistant integration, and correction workflows.",
])
heading(doc, "4.8 Discussion", 2)
paras(doc, [
    "The results show that the prototype's principal contribution is controlled integration. Scheduling rules, role relationships, clinical states, and AI safety behavior are encoded on the server and covered by tests. The system does not prove that waiting time is reduced in practice because no real-hospital baseline, deployment population, or longitudinal operational study was conducted.",
])

chapter(doc, "CHAPTER FIVE: SUMMARY, CONCLUSION, AND RECOMMENDATIONS")
heading(doc, "5.1 Summary", 2)
paras(doc, [
    "CareSync evolved beyond a three-role browser-storage scheduling mock-up into a four-role, PostgreSQL-backed full-stack prototype. It now connects secure identity, hospital administration, deterministic availability, patient booking, assigned-nurse intake, doctor consultation, patient communication, reporting, and bounded AI support.",
    "The final design treats the database and server as authoritative. Patients manage their own appointments and view their own information; nurses record vitals and triage for assigned visits; doctors complete consultations for assigned patients; administrators manage organizational resources and communication. Artificial intelligence supports preliminary routing but does not exercise clinical authority.",
])
heading(doc, "5.2 Achievement of Objectives", 2)
table(doc, ["Objective area", "Outcome"], [
    ("Secure roles and sessions", "Achieved in the prototype through bcrypt, JWT, rotating refresh sessions, RBAC and relationship checks"),
    ("Schedule and slot management", "Achieved with recurring schedules, exceptions, generated slots and transactional capacity"),
    ("Patient self-service", "Achieved for registration, profile, appointment management, symptoms, notifications and read-only clinical views"),
    ("Clinical workflow", "Achieved from check-in to consultation with assigned-nurse and assigned-doctor enforcement"),
    ("AI decision support", "Achieved within a restricted emergency-first symptom-assessment scope"),
    ("Deployment and testing", "Achieved for local Docker composition and automated prototype verification"),
])
heading(doc, "5.3 Conclusion", 2)
paras(doc, [
    "The project demonstrates that a role-scoped hospital workflow and cautiously designed AI component can operate within one web platform. Its strongest evidence is implementation and automated behavioral verification, not claims of clinical effectiveness or operational improvement. Explicit authorization and state transitions protect the meaning of records—for example, a value shown as verified on My Vitals is entered by the actively assigned nurse, not by the patient.",
    "CareSync is therefore a credible academic software prototype and a foundation for further evaluation. Production use would require governance, external security review, privacy and retention policies, backup and recovery processes, clinical validation, usability study, monitoring, resilient infrastructure, and formal approval by the responsible healthcare institution.",
])
heading(doc, "5.4 Recommendations", 2)
bullets(doc, [
    "Conduct structured usability testing with patients, nurses, doctors, and administrators and report participant count, method, and results.",
    "Use a dedicated test database and CI pipeline to run all JavaScript and Python suites on every change.",
    "Add production-grade shared rate limiting, centralized logging, monitoring, encrypted backups, secret management, HTTPS, and disaster-recovery procedures.",
    "Evaluate accessibility, low-bandwidth performance, and offline-safe behavior without caching sensitive API responses.",
    "Validate notification email delivery with a configured provider and expose delivery status instead of implying guaranteed transmission.",
    "Evaluate the symptom model with representative, ethically governed clinical data before considering any broader use; preserve clinician oversight and emergency-first behavior.",
    "Add standards-based integration only after governance requirements are defined, with FHIR-based interoperability as a possible future direction.",
    "Extend scope carefully to laboratory, pharmacy, billing, insurance, and inpatient functions only as separate validated modules.",
])
heading(doc, "5.5 Future Work", 2)
paras(doc, [
    "Future work may include production deployment, multi-facility support, richer audit review, delivery-status tracking, accessibility improvements, external healthcare integration, and real-world scheduling evaluation. Any future predictive scheduling model should be compared against a deterministic baseline using defined metrics such as conflict rate, utilization, no-show rate, average waiting time, and fairness across patient groups.",
])

chapter(doc, "APPENDIX A: OPERATIONAL WORKFLOW SUMMARY")
table(doc, ["Step", "Responsible actor", "System control"], [
    ("Account creation", "Patient or Admin for staff", "Public signup cannot create a staff role"),
    ("Availability setup", "Admin", "Schedules, assignments, exceptions and generated real slots"),
    ("Appointment booking", "Patient", "Ownership, relationships and slot capacity checked transactionally"),
    ("Nurse assignment", "Admin", "Active appointment-level assignment"),
    ("Check-in", "Nurse/Admin", "CONFIRMED to CHECKED_IN only"),
    ("Vital entry", "Assigned nurse", "Server derives recorder and verified nurse source"),
    ("Triage", "Nurse", "Clinician-selected urgency; CHECKED_IN to TRIAGED"),
    ("Queue", "Nurse/system ordering", "TRIAGED to WAITING; deterministic urgency/time ordering"),
    ("Consultation", "Assigned doctor", "WAITING to IN_CONSULTATION to COMPLETED"),
    ("Patient review", "Patient", "Own nurse vitals and finalized consultation summary only"),
])
heading(doc, "Appendix B: Local Deployment", 1)
paras(doc, [
    "The Docker Compose deployment runs the web client on port 80 (also mapped to 5173), the Express API on port 5000, the private AI service, and PostgreSQL. Environment files supply database, JWT, CORS, and optional email configuration. The API health endpoint is /api/health.",
    "This document reflects repository behavior reviewed on 2 September 2026. Later code changes should be accompanied by corresponding documentation and test updates.",
])

doc.core_properties.title = "CareSync AI-Powered Hospital Management System – Revised Documentation"
doc.core_properties.subject = "Implementation-accurate academic project documentation"
doc.core_properties.author = "CareSync Project"
doc.core_properties.comments = "Revised against the implemented repository on 2 September 2026."
doc.save(OUT)
print(OUT)
